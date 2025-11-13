"""
Tests for the resume parser module with NLP capabilities.
"""
import os
import pytest
import spacy
from resume_parser.parser import ResumeParser

# Load spaCy model for testing
nlp = spacy.load("en_core_web_sm")

def test_extract_text_from_docx(tmp_path):
    """Test DOCX text extraction"""
    # Create a test DOCX file
    from docx import Document
    doc = Document()
    doc.add_paragraph('Test Resume')
    doc.add_paragraph('Education')
    doc.add_paragraph('Bachelor in Computer Science')
    doc_path = os.path.join(tmp_path, "test_resume.docx")
    doc.save(doc_path)
    
    # Parse the test file
    parser = ResumeParser()
    text = parser.extract_text_from_docx(doc_path)
    
    assert 'Test Resume' in text
    assert 'Education' in text
    assert 'Bachelor in Computer Science' in text

def test_parse_sections():
    """Test section identification"""
    parser = ResumeParser()
    
    # Simulate parsed text
    parser.text = """
    John Doe
    Email: john@example.com
    Phone: 123-456-7890
    
    Summary
    Experienced software developer
    
    Education
    BS in Computer Science
    
    Experience
    Software Engineer at Tech Corp
    
    Skills
    Python, JavaScript, SQL
    """
    
    sections = parser.parse_resume(None)  # None because we set text directly
    
    assert 'contact' in sections
    assert 'summary' in sections
    assert 'education' in sections
    assert 'experience' in sections
    assert 'skills' in sections
    
    # Sections may be returned either as plain text strings or as dicts with
    # a 'text' field (older/newer formats). Normalize for assertions.
    def _section_text(s):
        if isinstance(s, dict):
            return s.get('text', '')
        return s

    assert 'john@example.com' in _section_text(sections['contact']).lower()
    assert 'software developer' in _section_text(sections['summary']).lower()
    assert 'computer science' in _section_text(sections['education']).lower()
    assert 'software engineer' in _section_text(sections['experience']).lower()
    assert 'python' in _section_text(sections['skills']).lower()

def test_get_contact_info():
    """Test contact information extraction"""
    parser = ResumeParser()
    
    # Simulate contact section
    parser.sections = {
        'contact': """
        John Doe
        Email: john@example.com
        Phone: 123-456-7890
        LinkedIn: linkedin.com/in/johndoe
        GitHub: github.com/johndoe
        """
    }
    
    contact_info = parser.get_contact_info()
    
    assert contact_info['email'] == 'john@example.com'
    assert '123-456-7890' in contact_info['phone']
    assert 'linkedin.com/in/johndoe' in contact_info['linkedin']
    assert 'github.com/johndoe' in contact_info['github']

def test_get_skills():
    """Test skills extraction"""
    parser = ResumeParser()
    
    # Simulate skills section
    parser.sections = {
        'skills': """
        Programming: Python, JavaScript, Java
        Databases: SQL, MongoDB
        Tools: Git, Docker
        """
    }
    
    skills = parser.get_skills()
    
    # Check for categorized skills
    assert 'Programming: Python' in skills
    assert 'Programming: JavaScript' in skills
    assert 'Programming: Java' in skills
    assert 'Databases: SQL' in skills
    assert 'Databases: MongoDB' in skills
    assert 'Tools: Git' in skills
    assert 'Cloud: Docker' in skills  # Docker is auto-categorized as Cloud

def test_get_education():
    """Test education information extraction"""
    parser = ResumeParser()
    
    # Simulate education section
    parser.section_details = {
        'education': {
            'text': """
            Stanford University
            Bachelor of Science in Computer Science, Expected 2024
            GPA: 3.8/4.0
            
            Harvard University
            Master of Science in Data Science, 2020
            GPA: 3.9/4.0
            """,
            'bullet_points': []
        }
    }
    
    education = parser.get_education()
    
    # Should find two degrees
    assert len(education) == 2
    
    # Check Stanford degree
    stanford = [e for e in education if e['school'] and 'Stanford' in e['school']][0]
    assert stanford['degree'] == 'Bachelors'
    assert stanford['major'] == 'Computer Science'
    assert stanford['graduation'] == 'Expected 2024'
    assert stanford['gpa'] == '3.8'
    
    # Check Harvard degree
    harvard = [e for e in education if e['school'] and 'Harvard' in e['school']][0]
    assert harvard['degree'] == 'Masters'
    assert harvard['major'] == 'Data Science'
    assert harvard['graduation'] == '2020'
    assert harvard['gpa'] == '3.9'

def test_get_experience():
    """Test experience information extraction"""
    parser = ResumeParser()
    
    # Simulate experience section
    parser.section_details = {
        'experience': {
            'text': """
            Senior Software Engineer at Google, Mountain View, CA
            January 2020 - Present
            • Led development of machine learning infrastructure serving 100M+ users
            • Optimized data pipeline reducing processing time by 40%
            • Mentored 5 junior engineers and conducted technical interviews
            
            Software Engineer at Microsoft, Redmond, WA
            June 2018 - December 2019
            • Developed new features for Azure Cloud Services
            • Improved system reliability by implementing automated testing
            • Collaborated with cross-functional teams on service integration
            """,
            'bullet_points': []
        }
    }
    
    experience = parser.get_experience()
    
    # Should find two positions
    assert len(experience) == 2
    
    # Check Google position
    google = [e for e in experience if e['company'] and 'Google' in e['company']][0]
    assert google['title'] == 'Senior Software Engineer'
    assert google['start_date'] == 'January 2020'
    assert google['end_date'] == 'Present'
    assert google['location'] == 'Mountain View, CA'
    assert len(google['achievements']) == 3
    assert any('machine learning' in bullet for bullet in google['achievements'])
    
    # Check Microsoft position
    microsoft = [e for e in experience if e['company'] and 'Microsoft' in e['company']][0]
    assert microsoft['title'] == 'Software Engineer'
    assert microsoft['start_date'] == 'June 2018'
    assert microsoft['end_date'] == 'December 2019'
    assert microsoft['location'] == 'Redmond, WA'
    assert len(microsoft['achievements']) == 3
    assert any('Azure' in bullet for bullet in microsoft['achievements'])